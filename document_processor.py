import os
from pathlib import Path
from typing import List, Dict
import PyPDF2
from docx import Document
import tempfile
import platform
import csv
import io
import json

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._doc_processing_methods = self._check_available_methods()
    
    def _check_available_methods(self) -> List[str]:
        """Check which .doc processing methods are available"""
        available = []
        
        # Check textutil (macOS)
        if platform.system() == "Darwin":
            try:
                import subprocess
                subprocess.run(['textutil', '-help'], capture_output=True, timeout=5)
                available.append('textutil')
            except:
                pass
        
        # Check doc2docx
        try:
            from doc2docx import convert
            available.append('doc2docx')
        except ImportError:
            pass
        
        # Check Windows COM
        if platform.system() == "Windows":
            try:
                import win32com.client
                available.append('win32com')
            except ImportError:
                pass
        
        # Check antiword
        try:
            import subprocess
            subprocess.run(['antiword', '-h'], capture_output=True, timeout=5)
            available.append('antiword')
        except:
            pass
        
        # Check catdoc
        try:
            import subprocess
            subprocess.run(['catdoc', '-h'], capture_output=True, timeout=5)
            available.append('catdoc')
        except:
            pass
        
        return available
    
    def get_doc_processing_status(self) -> Dict[str, any]:
        """Get status of .doc file processing capabilities"""
        return {
            'available_methods': self._doc_processing_methods,
            'primary_method': self._doc_processing_methods[0] if self._doc_processing_methods else None,
            'platform': platform.system(),
            'recommendations': self._get_setup_recommendations()
        }
    
    def _get_setup_recommendations(self) -> List[str]:
        """Get setup recommendations for better .doc processing"""
        recs = []
        
        if platform.system() == "Darwin":
            if 'textutil' not in self._doc_processing_methods:
                recs.append("textutil should be available by default on macOS")
            if 'doc2docx' not in self._doc_processing_methods:
                recs.append("Install doc2docx: pip install doc2docx")
        elif platform.system() == "Windows":
            if 'win32com' not in self._doc_processing_methods:
                recs.append("Install pywin32: pip install pywin32")
            if 'doc2docx' not in self._doc_processing_methods:
                recs.append("Install doc2docx: pip install doc2docx")
        else:  # Linux
            if 'antiword' not in self._doc_processing_methods:
                recs.append("Install antiword: sudo apt-get install antiword")
            if 'catdoc' not in self._doc_processing_methods:
                recs.append("Install catdoc: sudo apt-get install catdoc")
        
        return recs
    
    def process_file(self, file_path: str) -> List[Dict]:
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                text = self._extract_pdf_text(file_path)
            elif extension == '.docx':
                text = self._extract_docx_text(file_path)
            elif extension == '.doc':
                text = self._extract_doc_text(file_path)
            elif extension == '.csv':
                text = self._extract_csv_text(file_path)
            elif extension == '.md':
                text = self._extract_markdown_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {extension}")
            
            # Handle case where extraction returned placeholder text
            if not text or text.strip() == "":
                text = f"[Unable to extract content from {file_path.name}]"
                
        except Exception as e:
            print(f"Error processing {file_path.name}: {e}")
            text = f"[Error processing {file_path.name}: {str(e)}]"
        
        chunks = self._split_text(text)
        
        # Check for Google Drive metadata
        gdrive_metadata = {}
        metadata_file = Path(str(file_path) + '.gdrive_metadata')
        if metadata_file.exists():
            with open(metadata_file, 'r') as f:
                gdrive_metadata = json.load(f)
        
        documents = []
        for i, chunk in enumerate(chunks):
            metadata = {
                'source': str(file_path),
                'filename': file_path.name,
                'chunk_id': i,
                'file_type': extension[1:]
            }
            
            # Add Google Drive metadata if available
            if gdrive_metadata:
                metadata['gdrive_id'] = gdrive_metadata.get('gdrive_id')
                metadata['gdrive_link'] = gdrive_metadata.get('gdrive_link')
            
            documents.append({
                'content': chunk,
                'metadata': metadata
            })
        
        return documents
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        return text
    
    def _extract_docx_text(self, file_path: Path) -> str:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
        
        return text
    
    def _extract_doc_text(self, file_path: Path) -> str:
        """Extract text from .doc files using multiple fallback methods"""
        
        # Method 1: macOS textutil (most reliable on Mac)
        if platform.system() == "Darwin":
            try:
                import subprocess
                result = subprocess.run(
                    ['textutil', '-stdout', '-convert', 'txt', str(file_path)],
                    capture_output=True,
                    text=True,
                    check=True,
                    timeout=30
                )
                if result.stdout and result.stdout.strip():
                    return result.stdout
            except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired) as e:
                print(f"textutil failed: {e}")
        
        # Method 2: Windows COM automation
        if platform.system() == "Windows":
            try:
                import win32com.client
                import pythoncom
                pythoncom.CoInitialize()
                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(str(file_path))
                    text = doc.Range().Text
                    doc.Close()
                    word.Quit()
                    if text and text.strip():
                        return text
                finally:
                    pythoncom.CoUninitialize()
            except Exception as e:
                print(f"Windows COM failed: {e}")
        
        # Method 3: doc2docx conversion
        try:
            from doc2docx import convert
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp_path = tmp.name
            
            try:
                convert(str(file_path), tmp_path)
                if os.path.exists(tmp_path) and os.path.getsize(tmp_path) > 0:
                    text = self._extract_docx_text(Path(tmp_path))
                    if text and text.strip():
                        return text
            finally:
                # Always clean up temp file
                try:
                    os.unlink(tmp_path)
                except:
                    pass
                    
        except Exception as e:
            print(f"doc2docx conversion failed: {e}")
        
        # Method 4: Try python-docx anyway (sometimes works with .doc)
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            if text and text.strip():
                return text
        except Exception as e:
            print(f"python-docx fallback failed: {e}")
        
        # Method 5: Try antiword (if available on Linux/Mac)
        if platform.system() in ["Linux", "Darwin"]:
            try:
                import subprocess
                result = subprocess.run(
                    ['antiword', str(file_path)],
                    capture_output=True,
                    text=True,
                    check=True,
                    timeout=30
                )
                if result.stdout and result.stdout.strip():
                    return result.stdout
            except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
                pass
        
        # Method 6: Try catdoc (if available on Linux/Mac)
        if platform.system() in ["Linux", "Darwin"]:
            try:
                import subprocess
                result = subprocess.run(
                    ['catdoc', str(file_path)],
                    capture_output=True,
                    text=True,
                    check=True,
                    timeout=30
                )
                if result.stdout and result.stdout.strip():
                    return result.stdout
            except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
                pass
        
        # Method 7: Simple binary text extraction (last resort)
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                # Try to extract readable text from binary data
                text = ""
                for byte in raw_data:
                    if 32 <= byte <= 126:  # Printable ASCII
                        text += chr(byte)
                    elif byte in [10, 13]:  # Newlines
                        text += "\n"
                    else:
                        text += " "
                
                # Clean up the extracted text
                cleaned_text = " ".join(text.split())
                if len(cleaned_text) > 50:  # Only if we got substantial text
                    return cleaned_text
        except Exception as e:
            print(f"Binary extraction failed: {e}")
        
        # If all methods fail, provide helpful error message
        available_methods = ", ".join(self._doc_processing_methods) if self._doc_processing_methods else "none"
        recommendations = "; ".join(self._get_setup_recommendations())
        
        error_msg = f"""
Unable to read .doc file: {file_path.name}
Available methods: {available_methods}
Platform: {platform.system()}

Suggestions: {recommendations}
Alternative: Convert the file to .docx format manually.
"""
        print(error_msg)
        
        # Return a placeholder instead of raising an error
        return f"[Unable to read .doc file: {file_path.name}. Available methods: {available_methods}. Please convert to .docx format or check setup.]"
    
    def _extract_csv_text(self, file_path: Path) -> str:
        text = ""
        try:
            with open(file_path, 'r', encoding='utf-8', newline='') as csvfile:
                # Try to detect delimiter
                sample = csvfile.read(1024)
                csvfile.seek(0)
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
                
                reader = csv.reader(csvfile, delimiter=delimiter)
                for row_num, row in enumerate(reader):
                    if row_num == 0:
                        # Header row
                        text += "Headers: " + " | ".join(row) + "\n"
                    else:
                        text += "Row " + str(row_num) + ": " + " | ".join(row) + "\n"
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1', newline='') as csvfile:
                reader = csv.reader(csvfile)
                for row_num, row in enumerate(reader):
                    if row_num == 0:
                        text += "Headers: " + " | ".join(row) + "\n"
                    else:
                        text += "Row " + str(row_num) + ": " + " | ".join(row) + "\n"
        return text
    
    def _extract_markdown_text(self, file_path: Path) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _split_text(self, text: str) -> List[str]:
        chunks = []
        sentences = text.replace('\n', ' ').split('. ')
        
        current_chunk = ""
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            if len(current_chunk) + len(sentence) + 1 <= self.chunk_size:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        overlapped_chunks = []
        for i in range(len(chunks)):
            if i == 0:
                overlapped_chunks.append(chunks[i])
            else:
                overlap_text = chunks[i-1][-self.chunk_overlap:] if len(chunks[i-1]) > self.chunk_overlap else chunks[i-1]
                overlapped_chunks.append(overlap_text + " " + chunks[i])
        
        return overlapped_chunks